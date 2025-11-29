"""Document parsing utilities for various file formats"""

import re
from pathlib import Path

import markdown


class DocumentParser:
    """Parse documents from various file formats"""

    @staticmethod
    def parse_text(content: str) -> str:
        """Parse plain text content (already a string)"""
        return content.strip()

    @staticmethod
    def parse_markdown(content: str) -> str:
        """
        Parse markdown content and convert to plain text
        Preserves structure while removing markdown syntax
        """
        # Convert markdown to HTML first
        html = markdown.markdown(content)

        # Remove HTML tags
        text = re.sub(r"<[^>]+>", "", html)

        # Decode common HTML entities
        text = text.replace("&amp;", "&")
        text = text.replace("&lt;", "<")
        text = text.replace("&gt;", ">")
        text = text.replace("&quot;", '"')
        text = text.replace("&#39;", "'")

        return text.strip()

    @staticmethod
    def parse_pdf(file_path: str) -> str:
        """
        Extract text from PDF file

        Args:
            file_path: Path to PDF file

        Returns:
            Extracted text content
        """
        try:
            from pypdf import PdfReader
        except ImportError:
            raise ImportError("pypdf is required for PDF parsing. Install with: pip install pypdf")

        reader = PdfReader(file_path)
        text_parts = []

        for page in reader.pages:
            text = page.extract_text()
            if text:
                text_parts.append(text)

        return "\n\n".join(text_parts).strip()

    @staticmethod
    def parse_docx(file_path: str) -> str:
        """
        Extract text from DOCX file

        Args:
            file_path: Path to DOCX file

        Returns:
            Extracted text content
        """
        try:
            from docx import Document
        except ImportError:
            raise ImportError("python-docx is required for DOCX parsing. Install with: pip install python-docx")

        doc = Document(file_path)
        text_parts = []

        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                text_parts.append(paragraph.text)

        # Also extract text from tables
        for table in doc.tables:
            for row in table.rows:
                row_text = []
                for cell in row.cells:
                    if cell.text.strip():
                        row_text.append(cell.text)
                if row_text:
                    text_parts.append(" | ".join(row_text))

        return "\n\n".join(text_parts).strip()

    @staticmethod
    def detect_format(file_path: str) -> str:
        """
        Detect document format from file extension

        Args:
            file_path: Path to file

        Returns:
            Format identifier: 'pdf', 'docx', 'markdown', 'text'
        """
        path = Path(file_path)
        extension = path.suffix.lower()

        format_map = {
            ".pdf": "pdf",
            ".docx": "docx",
            ".doc": "docx",
            ".md": "markdown",
            ".markdown": "markdown",
            ".txt": "text",
        }

        return format_map.get(extension, "text")

    @classmethod
    def parse_file(cls, file_path: str, format: str | None = None) -> str:
        """
        Parse a document file automatically detecting format

        Args:
            file_path: Path to file
            format: Optional format override ('pdf', 'docx', 'markdown', 'text')

        Returns:
            Parsed text content
        """
        if format is None:
            format = cls.detect_format(file_path)

        if format == "pdf":
            return cls.parse_pdf(file_path)
        elif format == "docx":
            return cls.parse_docx(file_path)
        elif format == "markdown":
            with open(file_path, encoding="utf-8") as f:
                return cls.parse_markdown(f.read())
        else:  # text or unknown
            with open(file_path, encoding="utf-8") as f:
                return cls.parse_text(f.read())

    @staticmethod
    def clean_text(text: str) -> str:
        """
        Clean and normalize text

        Args:
            text: Raw text content

        Returns:
            Cleaned text
        """
        # Remove excessive whitespace
        text = re.sub(r"\s+", " ", text)

        # Remove excessive newlines (keep paragraph breaks)
        text = re.sub(r"\n{3,}", "\n\n", text)

        # Trim
        text = text.strip()

        return text
